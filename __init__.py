import azure.functions as func
import requests, logging
from get_secret import get_secret
from blob_connection import get_blob_url, saveblob
from azure.storage.blob import BlobClient
from docx import Document
import io
from pathlib import Path
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import time
import json

bp = func.Blueprint() 

@bp.function_name("TranscriptietoolWebApp")
@bp.blob_trigger(
    arg_name="inblob",
    path="hhh-inputfolder/{name}",
    connection="AzureWebJobsStorage"
)
def TranscriptietoolWebApp(inblob: func.InputStream):
    try:
        ## init variable and basic headers
        container_name = 'hhh-inputfolder'
        headers = {"authorization": get_secret('APIKEYAssemblyAI'), "content-type": "application/json"}
        fileName = inblob.name.split("/", 1)
        logging.info(f"init variable done")
        
        ## Upload audio file to assembly
        sas_url = get_blob_url(container_name=container_name, blob_name=fileName[1])
        
        # Create a client and download the blob into memory
        blob_client = BlobClient.from_blob_url(sas_url)
        # **Metadata ophalen**
        props = blob_client.get_blob_properties()
        metadata = props.metadata  # is een dict van str->str
        
        instructions = metadata.get('instructions') or ''
        blob_type = metadata.get('type') or 'transcript'
        tenantId = metadata.get('tenantId') or ''
        
        prompt = f"maak een {blob_type} {instructions}"
        
        stream = blob_client.download_blob()
        audio_bytes = stream.readall()
        
        # make request
        r = requests.post("https://api.eu.assemblyai.com/v2/upload",
                        headers=headers, data=audio_bytes, timeout=30)
        r.raise_for_status()
            
        # get upload URL
        data = r.json()
        try:
            upload_url = data["upload_url"]
        except Exception as ex:
            logging.error("No upload URL") 
            return None
        
        logging.info("Uploaded audio file")      
        
        ## Transcription start 
        payload = {"audio_url": upload_url, "speaker_labels": "true", "speakers_expected": 3, "language_detection": "true"}

        r = requests.post("https://api.eu.assemblyai.com/v2/transcript",
                        headers=headers, json=payload, timeout=30)
        r.raise_for_status()
        data = r.json()
        transcriptid = data['id']
        
        logging.info("Started Transcription")
        
        ## Pulse Transcription
        response = wait_for_transcript(request_id=transcriptid, headers=headers)  
        logging.info("Finished Transcription")
                
        if (blob_type != "transcript"):
            ## Get LeMUR  
            payload = {
                "final_model": "anthropic/claude-3-5-sonnet",
                "prompt": prompt,
                "max_output_size": 4000,
                "temperature": 0,
                "transcript_ids": [transcriptid]
            }
            r = requests.post("https://api.eu.assemblyai.com/lemur/v3/generate/task",
                            headers=headers, json=payload, timeout=30)
            r.raise_for_status()
            data = r.json() 
            request_id = data['request_id']
            response = data['response']
            logging.info("Finished LeMUR")  

            ## Delete LeMUR
            r = requests.delete(f"https://api.eu.assemblyai.com/lemur/v3/{request_id}",
                            headers=headers, timeout=30)
            if r.status_code == 200:
                logging.info(f"LeMUR is correctly deleted, {request_id}")
            else:
                logging.error(f"LeMUR not deleted, ID:{request_id}")
                
        ## Delete Transcription
        r = requests.delete(f"https://api.eu.assemblyai.com/v2/transcript/{transcriptid}",
                        headers=headers, timeout=30)
        if r.status_code == 200:
            logging.info(f"Transcript is correctly deleted, {transcriptid}")
        else:
            logging.error(f"Transcript not deleted, ID:{transcriptid}")
                               
        buf = generate_word_document(response)
        
        p = Path(inblob.name)
        # Derive just the filename (without the container prefix)
        file_stem = p.stem
        
        folder_id = p.parent.name
        
        saveblob(file_name=f"{folder_id}/{file_stem}.docx", data=buf.getvalue(), container_name="hhh-outputfolder")
        logging.info(f"File uploaded and function done!")
        
    except requests.HTTPError as e:
        logging.error(f"Request failed: {e}")   # includes status and reason
        
    except Exception as ex :
        logging.error(ex)
        
def wait_for_transcript(
    request_id: str,
    headers: dict,
    poll_interval: float = 3.0,   # seconds between polls
    timeout: float = 3600.0        # overall timeout in seconds
) -> dict:
    """
    Polls the AssemblyAI transcript endpoint until status == 'completed',
    or until timeout. Treats any non-200 response as 'not ready yet'.
    """
    start = time.monotonic()
    url   = f"https://api.eu.assemblyai.com/v2/transcript/{request_id}"

    while True:
        resp = requests.get(url, headers=headers, timeout=10)

        # If it's a 200 OK, parse and check the JSON status
        if resp.status_code == 200:
            data   = resp.json()
            status = data.get("status", "").lower()
            if status == "completed":
                return data
            if status == "error":
                raise RuntimeError(f"Transcription failed: {data}")
        else:
            # Any other status (400, 202, etc) → still pending
            # you can log resp.status_code if you like
            pass

        # Check timeout
        elapsed = time.monotonic() - start
        if elapsed >= timeout:
            raise TimeoutError(f"Polling timed out after {timeout:.1f}s")
        logging.info("Polling")
        time.sleep(poll_interval)


def generate_word_document(data: dict):
    """
    Genereert een Word-document op basis van de JSON-input.
   
    Ondersteunde structuren:
    1. Als er een 'variables' key bestaat, zoekt hij naar het object met "name": "transcriptBodyComplete".
    2. Anders gaat hij ervan uit dat de hele JSON de transcriptiegegevens bevat.
   
    Verwacht dat de transcriptiegegevens een "words" veld bevatten met een lijst van woorden.
   
    Het document bevat:
    - Header met transcriptiedatum en logo.
    - Indien aanwezig: een samenvatting (onder de key "samenvatting") met een header.
    - De transcriptie: voor elke spreker wordt een regel toegevoegd in het formaat:
         Speaker X: tekst...
    Retourneert een base64-string van het gegenereerde document.
    """
    logging.info("Start met genereren van Word-document met aangepaste JSON-structuur")
 
    # Bepaal of we de transcript data uit een 'variables'-object moeten halen
    if "variables" in data:
        variables = data.get("variables", [])
        transcript_data = None
        for var in variables:
            if var.get("name") == "transcriptBodyComplete":
                transcript_data = var.get("value", {})
                break
        if transcript_data is None:
            raise ValueError("Geen transcript data gevonden in 'variables'.")
    else:
        transcript_data = data
 
    # Haal de woordenlijst op
    words = transcript_data.get("words", [])
    if not words:
        raise ValueError("Geen 'words' veld gevonden in de transcript data.")
 
    # Groepeer de woorden per opeenvolgende spreker
    transcript_groups = []
    current_speaker = None
    current_texts = []
    for word in words:
        speaker = word.get("speaker")
        text = word.get("text", "")
        if speaker is None:
            continue
        if speaker != current_speaker:
            if current_texts:
                transcript_groups.append((current_speaker, " ".join(current_texts)))
            current_speaker = speaker
            current_texts = [text]
        else:
            current_texts.append(text)
    if current_texts:
        transcript_groups.append((current_speaker, " ".join(current_texts)))
 
    # Maak een mapping van originele spreker identifiers naar genummerde speakers (1, 2, ...)
    speaker_mapping = {}
    next_speaker_number = 1
    transcript_lines = []
    for speaker, group_text in transcript_groups:
        if speaker not in speaker_mapping:
            speaker_mapping[speaker] = next_speaker_number
            next_speaker_number += 1
        speaker_label = f"Speaker {speaker_mapping[speaker]}:"
        transcript_lines.append((speaker_label, group_text))
 
    # Maak een nieuw Word-document
    doc = Document()
 
    # --- Header maken ---
    section = doc.sections[0]
    header = section.header
 
    # Titel links in de header met transcriptiedatum
    current_date = datetime.now().strftime("%d-%m-%Y")
    title_paragraph = header.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = title_paragraph.add_run(f"Transcriptie {current_date}")
    run_title.bold = True
 
    # --- Eventuele samenvatting toevoegen ---
    # Controleer of er een 'samenvatting' key aanwezig is in de top-level JSON
    samenvatting_tekst = data.get("samenvatting")
    if samenvatting_tekst:
        samenvatting_header = doc.add_paragraph()
        run_summary_header = samenvatting_header.add_run("Korte samenvatting:")
        run_summary_header.bold = True
        doc.add_paragraph(samenvatting_tekst)
        doc.add_paragraph("")
 
    # --- Transcriptie toevoegen ---
    transcriptie_header = doc.add_paragraph()
    transcriptie_run = transcriptie_header.add_run("Transcriptie:")
    transcriptie_run.bold = True
 
    for speaker_label, text in transcript_lines:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{speaker_label} ")
        run_label.bold = True
        p.add_run(text)
 
    # Document opslaan in geheugen en omzetten naar base64
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io